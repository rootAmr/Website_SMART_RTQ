from flask import Flask, Flask, render_template, jsonify, request, send_file
import pandas as pd
import requests
from datetime import datetime
from data_endpoint import df, file_path  
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from io import BytesIO

app = Flask(__name__)

# Function to calculate age from birth date string
def calculate_age(birth_date_str):
    today = datetime.today()
    try:
        place, date_str = birth_date_str.split(', ')
        day, month_str, year = date_str.split(' ')
        months = {
            'Januari': 1, 'Februari': 2, 'Maret': 3, 'April': 4, 'Mei': 5, 'Juni': 6,
            'Juli': 7, 'Agustus': 8, 'September': 9, 'Oktober': 10, 'November': 11, 'Desember': 12
        }
        month = months[month_str]
        birth_date = datetime(year=int(year), month=month, day=int(day))
        age = today.year - birth_date.year - ((today.month, today.day) < (birth_date.month, birth_date.day))
        return age
    except Exception as e:
        print(f"Error calculating age: {e}")
        return None

# Function to determine age category
def determine_age_category(age):
    if age is None:
        return None
    if 0 <= age <= 6:
        return '0-6'
    elif 7 <= age <= 9:
        return '7-9'
    elif 10 <= age <= 13:
        return '10-13'
    elif 14 <= age <= 16:
        return '14-16'
    elif 17 <= age <= 19:
        return '17-19'
    elif 20 <= age <= 30:
        return '20-30'
    elif 31 <= age <= 40:
        return '31-40'
    elif 41 <= age <= 50:
        return '41-50'
    elif 51 <= age <= 60:
        return '51-60'
    elif age >= 61:
        return '61 keatas'
    return None

# Apply the functions to the DataFrame
df['Age'] = df['T.LAHIR'].apply(calculate_age)
df['rentang_umur'] = df['Age'].apply(determine_age_category)

# Save the changes to the Excel file
df.to_excel(file_path, index=False)

# Route to render the dashboard
@app.route('/')
def index():
    return render_template('index.html')

@app.route('/dashboard')
def dashboard():
    return render_template('dashboard.html')

@app.route('/warga')
def warga():
    return render_template('warga.html')

@app.route('/laporan')
def laporan():
    return render_template('file_laporan.html')

@app.route('/tambah/warga')
def tambahwarga():
    return render_template('tambah_warga.html')

@app.route('/edit/warga')
def editwarga():
    return render_template('edit_warga.html')

@app.route('/edit/formulir/warga')
def formulireditwarga():
    resident_id = request.args.get('nik')
    
    # Assuming `df` is your DataFrame with resident data
    matching_resident = df[df['NIK'].astype(str) == resident_id]

    if matching_resident.empty:
        # Handle case where no matching resident is found
        return "No resident found", 404
    
    # Convert the matching resident data to a dictionary
    resident_data = matching_resident.to_dict(orient='records')[0]

    return render_template('formulir_warga_edit.html', resident=resident_data)

@app.route('/agerange', methods=['GET'])
def get_age_range_counts():
    status = request.args.get('status', 'PERMANEN').upper()

    # Standardize the 'JENIS KELAMIN' column
    df['JENIS KELAMIN'] = df['JENIS KELAMIN'].str.strip().str.lower()

    # Define the age ranges
    age_ranges = ['0-6', '7-9', '10-13', '14-16', '17-19', '20-30', '31-40', '41-50', '51-60', '61 keatas']

    # Initialize dictionaries to store counts
    counts = {'male': {}, 'female': {}, 'total': {}}

    # Filter the DataFrame for the selected status
    df_status = df[df['Status Kependudukan'] == status]

    # Count occurrences for each age range and gender
    age_range_counts_male = df_status[df_status['JENIS KELAMIN'] == 'laki-laki']['rentang_umur'].value_counts().reindex(age_ranges, fill_value=0).astype(int)
    age_range_counts_female = df_status[df_status['JENIS KELAMIN'] == 'perempuan']['rentang_umur'].value_counts().reindex(age_ranges, fill_value=0).astype(int)

    # Combine male and female counts
    age_range_counts_total = age_range_counts_male + age_range_counts_female

    # Store the results in the dictionary
    for age_range in age_ranges:
        counts['male'][age_range] = int(age_range_counts_male[age_range])
        counts['female'][age_range] = int(age_range_counts_female[age_range])
        counts['total'][age_range] = int(age_range_counts_total[age_range])

    # Prepare the results in the desired format
    result = {}
    for age_range in age_ranges:
        result[f'L-{age_range}'] = counts['male'][age_range]
        result[f'P-{age_range}'] = counts['female'][age_range]
        result[f'T-{age_range}'] = counts['total'][age_range]

    return jsonify(result)

# Function to get the month in Indonesian
def get_indonesian_month(month_number):
    months = [
        "Januari", "Februari", "Maret", "April", "Mei", "Juni",
        "Juli", "Agustus", "September", "Oktober", "November", "Desember"
    ]
    return months[month_number - 1]

@app.route('/generate-document', methods=['GET'])
def generate_document():
    status = request.args.get('status', 'PERMANEN').upper()  # Get status from query params, default to 'PERMANEN'
    api_url = f'http://localhost:5000/agerange?status={status}'

    # Fetch data from the API
    try:
        response = requests.get(api_url)
        response.raise_for_status()  # Check if request was successful
        data = response.json()
    except requests.RequestException as e:
        return jsonify({'error': f'Failed to fetch data: {e}'}), 500

    if not data:
        return jsonify({'error': 'No data available to generate the document.'}), 500

    # Load the template document
    template_path = r'D:\DASHBOARD_RTQ\Dokumen_Template.docx'
    try:
        doc = Document(template_path)
    except Exception as e:
        return jsonify({'error': f'Failed to load template: {e}'}), 500

    # Get the current month and year
    current_date = datetime.now()
    bulan = get_indonesian_month(current_date.month)
    tahun = current_date.year

    # Add month and year to data for replacement
    data['Bulan'] = bulan
    data['Tahun'] = tahun

    # Replace placeholders with data, center text, and set font to Times New Roman
    for key, value in data.items():
        placeholder = f'{{{{{key}}}}}'
        for paragraph in doc.paragraphs:
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, str(value))
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
                # Center text except for 'Tahun' and 'Bulan'
                if key not in ['Bulan', 'Tahun']:
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if placeholder in cell.text:
                        cell.text = cell.text.replace(placeholder, str(value))
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = 'Times New Roman'
                        # Center text except for 'Tahun' and 'Bulan'
                        if key not in ['Bulan', 'Tahun']:
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Format the filename
    filename = f'Data_Warga_{status}_{current_date.strftime("%d %B %Y")}.docx'

    # Save the document to a BytesIO object
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    # Send the document with a 200 OK status and success message
    response = send_file(buffer, as_attachment=True, download_name=filename, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response.headers['Custom-Message'] = 'Document generated successfully'
    return response, 200

# API route to get the count of gender
@app.route('/api/count/gender')
def count_gender():
    # Clean the 'JENIS KELAMIN' column
    gender_series = df['JENIS KELAMIN'].str.strip().str.lower()

    # Calculate the count for male and female
    male_count = gender_series[gender_series == 'laki-laki'].shape[0]
    female_count = gender_series[gender_series == 'perempuan'].shape[0]

    return jsonify({'male_count': male_count, 'female_count': female_count})

# API route to get the count of resident statuses
@app.route('/api/count/status')
def count_status():
    # Get the value counts of the 'Status Kependudukan' column
    status_counts = df['Status Kependudukan'].value_counts().to_dict()
    return jsonify(status_counts)

# API route to get the count of age ranges for permanent residents
@app.route('/api/count/age_ranges/permanent')
def count_age_ranges_permanent():
    # Filter the DataFrame for permanent residents and create a copy
    df_permanent = df[df['Status Kependudukan'] == 'PERMANEN'].copy()

    # Clean the 'JENIS KELAMIN' column
    df_permanent['JENIS KELAMIN'] = df_permanent['JENIS KELAMIN'].str.strip().str.lower()

    # Define the age ranges
    age_ranges = ['0-6', '7-9', '10-13', '14-16', '17-19', 
                  '20-30', '31-40', '41-50', '51-60', '61 keatas']

    # Initialize a dictionary to store the counts
    age_range_counts_permanent = {age_range: {'L': 0, 'P': 0, 'Total': 0} for age_range in age_ranges}

    # Iterate over each age range to calculate the counts
    for age_range in age_ranges:
        male_count = df_permanent[(df_permanent['rentang_umur'] == age_range) & 
                                  (df_permanent['JENIS KELAMIN'] == 'laki-laki')].shape[0]
        female_count = df_permanent[(df_permanent['rentang_umur'] == age_range) & 
                                    (df_permanent['JENIS KELAMIN'] == 'perempuan')].shape[0]
        total_count = male_count + female_count

        # Store the counts in the dictionary
        age_range_counts_permanent[age_range]['L'] = male_count
        age_range_counts_permanent[age_range]['P'] = female_count
        age_range_counts_permanent[age_range]['Total'] = total_count

    return jsonify(age_range_counts_permanent)


# API route to get the total count of residents
@app.route('/api/count/total')
def count_total_residents():
    # Calculate the total number of residents
    total_residents = df.shape[0]
    return jsonify({'total_residents': total_residents})

# API route to retrieve resident data
@app.route('/api/data/residents', methods=['GET'])
def get_residents_data():
    # Select the relevant columns
    columns = ['NAMA', 'Status Kependudukan', 'NIK', 'T.LAHIR', 'JENIS KELAMIN', 'PEKERJAAN', 'ALAMAT', 'AGAMA', 'Age', 'rentang_umur']
    
    # Create a DataFrame with only the selected columns
    residents_data = df[columns]

    # Convert the DataFrame to a dictionary
    residents_dict = residents_data.to_dict(orient='records')

    return jsonify(residents_dict)

@app.route('/api/add/resident', methods=['POST'])
def add_resident():
    # Get the data from the request
    data = request.json

    # Ensure all fields are present and not empty
    required_fields = ["NAMA", "Status Kependudukan", "NIK", "T.LAHIR", "JENIS KELAMIN", 
                       "PEKERJAAN", "ALAMAT", "AGAMA", "Age", "rentang_umur"]
    
    missing_fields = [field for field in required_fields if field not in data or not data[field]]
    if missing_fields:
        return jsonify({'error': f'Missing or empty fields: {", ".join(missing_fields)}'}), 400

    # Check if the NIK already exists in the DataFrame
    global df
    new_nik = data['NIK']
    existing_record = df[df['NIK'].astype(str) == new_nik]

    if not existing_record.empty:
        existing_name = existing_record['NAMA'].values[0]
        return jsonify({'error': f'NIK {new_nik} Sudah di Gunakan Oleh {existing_name}'}), 400

    # Convert text fields to Proper Case
    def to_proper_case(text):
        return text.title() if text else text

    # Apply Proper Case conversion
    data["NAMA"] = to_proper_case(data.get("NAMA", ""))
    data["T.LAHIR"] = to_proper_case(data.get("T.LAHIR", ""))
    data["PEKERJAAN"] = to_proper_case(data.get("PEKERJAAN", ""))
    data["ALAMAT"] = to_proper_case(data.get("ALAMAT", ""))

    # Create a DataFrame from the incoming data
    new_data_df = pd.DataFrame([data])

    # Append the new data to the existing DataFrame
    df = pd.concat([df, new_data_df], ignore_index=True)

    # Save the updated DataFrame back to the Excel file
    try:
        df.to_excel('./data/Data_Warga17_FIX.xlsx', index=False)
    except Exception as e:
        return jsonify({'error': f'Failed to save data: {str(e)}'}), 500

    return jsonify({'message': 'Data Berhasil Ditambahkan'}), 201

@app.route('/api/search/resident', methods=['GET'])
def search_resident():
    # Get the search query from the request arguments
    nik_query = request.args.get('nik', '').strip()

    # Check if the NIK query is provided
    if not nik_query:
        return jsonify({'error': 'NIK query parameter is required'}), 400

    # Filter the DataFrame for the matching NIK
    matching_resident = df[df['NIK'].astype(str) == nik_query]

    # Check if no matching resident is found
    if matching_resident.empty:
        return jsonify({'error': 'No resident found with the given NIK'}), 404

    # Select the relevant columns to return
    columns = ['NAMA', 'Status Kependudukan', 'NIK', 'T.LAHIR', 'JENIS KELAMIN', 
               'PEKERJAAN', 'ALAMAT', 'AGAMA', 'Age', 'rentang_umur']

    # Convert the matching record to a dictionary
    resident_dict = matching_resident[columns].to_dict(orient='records')

    return jsonify(resident_dict), 200

@app.route('/api/update/resident', methods=['PUT'])
def update_resident():
    global df  # Declare the global variable at the beginning of the function
    
    # Get the old NIK from the query string
    old_nik_query = request.args.get('nik', '').strip()
    if not old_nik_query:
        return jsonify({'error': 'NIK query parameter is required'}), 400

    # Get the updated data from the request body
    data = request.json

    # Ensure that all required fields are present and not empty
    required_fields = ['NIK', 'NAMA', 'Status Kependudukan', 'T.LAHIR', 'JENIS KELAMIN', 
                       'PEKERJAAN', 'ALAMAT', 'AGAMA', 'Age', 'rentang_umur']
    
    missing_fields = [field for field in required_fields if field not in data or not data[field]]
    if missing_fields:
        return jsonify({'error': f'Missing or empty fields: {", ".join(missing_fields)}'}), 400

    # Convert text fields to Proper Case
    def to_proper_case(text):
        return text.title() if text else text

    # Apply Proper Case conversion
    data["NAMA"] = to_proper_case(data.get("NAMA", ""))
    data["T.LAHIR"] = to_proper_case(data.get("T.LAHIR", ""))
    data["PEKERJAAN"] = to_proper_case(data.get("PEKERJAAN", ""))
    data["ALAMAT"] = to_proper_case(data.get("ALAMAT", ""))

    # Check if the new NIK already exists in the DataFrame (excluding the old one)
    new_nik = data['NIK']
    existing_record = df[(df['NIK'].astype(str) == data['NIK']) & (df['NIK'].astype(str) != old_nik_query)]

    if not existing_record.empty:
        existing_name = existing_record['NAMA'].values[0]
        return jsonify({'error': f'NIK {data["NIK"]} Sudah di Gunakan Oleh {existing_name}. Pastikan NIK yang anda Ketik Sudah Benar'}), 400

    # Find the index of the resident with the specified old NIK
    resident_index = df[df['NIK'].astype(str) == old_nik_query].index

    if resident_index.empty:
        return jsonify({'error': 'No resident found with the given NIK'}), 404

    # Update the resident's information
    df.loc[resident_index, 'NIK'] = new_nik
    df.loc[resident_index, 'NAMA'] = data['NAMA']
    df.loc[resident_index, 'Status Kependudukan'] = data['Status Kependudukan']
    df.loc[resident_index, 'T.LAHIR'] = data['T.LAHIR']
    df.loc[resident_index, 'JENIS KELAMIN'] = data['JENIS KELAMIN']
    df.loc[resident_index, 'PEKERJAAN'] = data['PEKERJAAN']
    df.loc[resident_index, 'ALAMAT'] = data['ALAMAT']
    df.loc[resident_index, 'AGAMA'] = data['AGAMA']
    df.loc[resident_index, 'Age'] = data['Age']
    df.loc[resident_index, 'rentang_umur'] = data['rentang_umur']

    # Save the updated DataFrame back to the Excel file
    try:
        df.to_excel(r'./data/Data_Warga17_FIX.xlsx', index=False)
    except Exception as e:
        return jsonify({'error': f'Failed to save data: {str(e)}'}), 500

    return jsonify({'message': 'Data Berhasil di Update'}), 200

@app.route('/api/delete/resident', methods=['DELETE'])
def delete_resident():
    global df  # Declare the global variable at the beginning of the function
    
    # Get the NIK from the query string
    nik_query = request.args.get('nik', '').strip()
    if not nik_query:
        return jsonify({'error': 'NIK query parameter is required'}), 400

    # Find the index of the resident with the specified NIK
    resident_index = df[df['NIK'].astype(str) == nik_query].index

    if resident_index.empty:
        return jsonify({'error': 'No resident found with the given NIK'}), 404

    # Drop the resident's row from the DataFrame
    df = df.drop(resident_index)

    # Save the updated DataFrame back to the Excel file
    try:
        df.to_excel(file_path, index=False)
    except Exception as e:
        return jsonify({'error': f'Failed to save data: {str(e)}'}), 500

    return jsonify({'message': 'Data Berhasil di Hapus'}), 200


if __name__ == '__main__':
    app.run(debug=True)
